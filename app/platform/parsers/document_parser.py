"""
文档解析器
"""
from typing import Optional, Dict, Any, List
from dataclasses import dataclass
from pathlib import Path
from abc import ABC, abstractmethod
import mimetypes


@dataclass
class ParseResult:
    """解析结果"""
    title: str
    content: str
    metadata: Dict[str, Any]
    chunks: Optional[List[str]] = None
    error: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            "title": self.title,
            "content": self.content,
            "metadata": self.metadata,
            "chunk_count": len(self.chunks) if self.chunks else 0,
            "error": self.error,
        }


class DocumentParser(ABC):
    """文档解析器抽象基类"""
    
    @abstractmethod
    async def parse(self, file_path: str) -> ParseResult:
        """解析文档"""
        pass
    
    @abstractmethod
    def supports(self, file_path: str) -> bool:
        """检查是否支持该文件类型"""
        pass


class PDFParser(DocumentParser):
    """PDF解析器"""
    
    async def parse(self, file_path: str) -> ParseResult:
        """解析PDF"""
        try:
            try:
                import PyPDF2
            except ImportError:
                try:
                    import pypdf as PyPDF2
                except ImportError:
                    return ParseResult(
                        title="",
                        content="",
                        metadata={},
                        error="PyPDF2 or pypdf is not installed",
                    )
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                title = Path(file_path).stem
                content = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    content += page.extract_text() + "\n"
                
                metadata = {
                    "file_path": file_path,
                    "file_type": "pdf",
                    "page_count": len(pdf_reader.pages),
                }
                
                return ParseResult(
                    title=title,
                    content=content,
                    metadata=metadata,
                )
        except Exception as e:
            return ParseResult(
                title="",
                content="",
                metadata={},
                error=f"Failed to parse PDF: {str(e)}",
            )
    
    def supports(self, file_path: str) -> bool:
        """检查是否支持PDF"""
        return file_path.lower().endswith('.pdf')


class DOCXParser(DocumentParser):
    """DOCX解析器"""
    
    async def parse(self, file_path: str) -> ParseResult:
        """解析DOCX"""
        try:
            try:
                from docx import Document
            except ImportError:
                return ParseResult(
                    title="",
                    content="",
                    metadata={},
                    error="python-docx is not installed",
                )
            
            doc = Document(file_path)
            title = doc.core_properties.title or Path(file_path).stem
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            metadata = {
                "file_path": file_path,
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
            }
            
            return ParseResult(
                title=title,
                content=content,
                metadata=metadata,
            )
        except Exception as e:
            return ParseResult(
                title="",
                content="",
                metadata={},
                error=f"Failed to parse DOCX: {str(e)}",
            )
    
    def supports(self, file_path: str) -> bool:
        """检查是否支持DOCX"""
        return file_path.lower().endswith(('.docx', '.doc'))


class TXTParser(DocumentParser):
    """文本解析器"""
    
    async def parse(self, file_path: str) -> ParseResult:
        """解析文本文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                return ParseResult(
                    title="",
                    content="",
                    metadata={},
                    error="Failed to decode file with any encoding",
                )
            
            title = Path(file_path).stem
            metadata = {
                "file_path": file_path,
                "file_type": "txt",
            }
            
            return ParseResult(
                title=title,
                content=content,
                metadata=metadata,
            )
        except Exception as e:
            return ParseResult(
                title="",
                content="",
                metadata={},
                error=f"Failed to parse text file: {str(e)}",
            )
    
    def supports(self, file_path: str) -> bool:
        """检查是否支持文本文件"""
        return file_path.lower().endswith(('.txt', '.md', '.markdown'))


class ExcelParser(DocumentParser):
    """Excel解析器"""
    
    async def parse(self, file_path: str) -> ParseResult:
        """解析Excel"""
        try:
            import pandas as pd
            
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            title = Path(file_path).stem
            content_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                content_parts.append(f"Sheet: {sheet_name}\n")
                content_parts.append(df.to_string())
                content_parts.append("\n\n")
            
            content = "\n".join(content_parts)
            
            metadata = {
                "file_path": file_path,
                "file_type": "excel",
                "sheet_count": len(excel_file.sheet_names),
                "sheet_names": excel_file.sheet_names,
            }
            
            return ParseResult(
                title=title,
                content=content,
                metadata=metadata,
            )
        except Exception as e:
            return ParseResult(
                title="",
                content="",
                metadata={},
                error=f"Failed to parse Excel: {str(e)}",
            )
    
    def supports(self, file_path: str) -> bool:
        """检查是否支持Excel"""
        return file_path.lower().endswith(('.xlsx', '.xls'))


class HTMLParser(DocumentParser):
    """HTML解析器"""
    
    async def parse(self, file_path: str) -> ParseResult:
        """解析HTML"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 提取标题
            title_tag = soup.find('title')
            title = title_tag.text if title_tag else Path(file_path).stem
            
            # 提取文本内容
            for script in soup(["script", "style"]):
                script.decompose()
            
            content = soup.get_text(separator='\n', strip=True)
            
            metadata = {
                "file_path": file_path,
                "file_type": "html",
            }
            
            return ParseResult(
                title=title,
                content=content,
                metadata=metadata,
            )
        except Exception as e:
            return ParseResult(
                title="",
                content="",
                metadata={},
                error=f"Failed to parse HTML: {str(e)}",
            )
    
    def supports(self, file_path: str) -> bool:
        """检查是否支持HTML"""
        return file_path.lower().endswith(('.html', '.htm'))


class JSONParser(DocumentParser):
    """JSON解析器"""
    
    async def parse(self, file_path: str) -> ParseResult:
        """解析JSON"""
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            title = Path(file_path).stem
            content = json.dumps(data, ensure_ascii=False, indent=2)
            
            metadata = {
                "file_path": file_path,
                "file_type": "json",
            }
            
            return ParseResult(
                title=title,
                content=content,
                metadata=metadata,
            )
        except Exception as e:
            return ParseResult(
                title="",
                content="",
                metadata={},
                error=f"Failed to parse JSON: {str(e)}",
            )
    
    def supports(self, file_path: str) -> bool:
        """检查是否支持JSON"""
        return file_path.lower().endswith('.json')

